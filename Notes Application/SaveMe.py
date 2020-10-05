#!pip install PySimpleGUI
#!pip install python-docx
import PySimpleGUI as sg
import sys
import os
import docx
import time
import pyautogui
def take_screenshot(path):
       
    named_tuple = time.localtime()
    folder_name=time.strftime("%d-%m-%Y",named_tuple)
    path=path+'/Screenshots on '+folder_name+'/'
    try:
        os.makedirs(path)
    except:
        pass
    
    named_tuple = time.localtime()
    time_string=time.strftime("%m-%d-%Y %H.%M.%S",named_tuple)

    myScreenshot = pyautogui.screenshot()
    myScreenshot.save(path+time_string+'.png')
    image_name=path+time_string+'.png'
    return(image_name,time_string)

def add_image_to_doc(mydoc,image,path,string):
    mydoc.add_picture(image, width=docx.shared.Inches(5.99), height=docx.shared.Inches(3.36))
    mydoc.save(path+"\\"+"mydoc"+'.docx')
        

def add_content_to_doc(mydoc,content,path):
    mydoc.add_paragraph(content)
    mydoc.save(path+"\\"+"mydoc"+'.docx')


sg.theme('SandyBeach') 

layout = [  [sg.Button('Take Screenshot',size=(20,2)),sg.Button('Take Notes',size=(20,2))],
            [sg.Text('Notes'), sg.Multiline('', key='notes',size=(40, 2))],
            [sg.Text('Session Name'), sg.InputText('', key='session',size=(20, 10)),sg.Button('Create Folder')],
            [sg.InputText(key="path", default_text=os.getcwd(), size=(40, 1)), sg.FolderBrowse()]
         ]

window = sg.Window('Save me',icon=r"C:\Users\vk00489900\Desktop\Python\Vamsi Scripts\Logo.ico",resizable=True).Layout(layout)

sg.popup("Save Me", """This is a notes taking application. You can save screenshots along with comments \
in a word document using this application without any hassle.\n \n 1.You may begin by selecting the path where the document is to be saved.\n
2.Enter the session name.(If you are using this for a particular service, Give a service name+date - eg :Client Reporting_15-09-20))\n
3.A folder with your session name is created. And your document is saved as mydoc.docx and your screenshots are saved in a separate folder inside it.\n
4.Click on Save Screenshot to save the screenshot of the windows.\n
5.Enter comments inside box and click on Take Notes Button to save it in word Document.\n
6.Congratulations on reading until here. You seem to have quite a lot of patience. Without a do, lets Jump in......""")
ss_taken=False
t_count=1
n_taken=False
i_count=1
create_button_count=0
og_path=os.getcwd()
try:
    while True:
        event, values = window.read()
        
        if event == "Create Folder":
            try:
                if create_button_count ==0:
                    val=values['path']+'/'+values['session']
                    os.makedirs(values['path']+'/'+values['session'])
                    os.chdir(values['path']+'/'+values['session'])
                    mydoc = docx.Document() 
                    mydoc.save(val+'/'+"mydoc"+'.docx')
                    session_value=values['session']
                    create_button_count+=1
                else:
                    sg.popup('Alert!!','Alert!! Folder you are trying to create is already Present. Try giving a new Folder Name.\nRecommended to give a date as postfix. eg: Client_Reporting_20-08-20. \nIf you want to append to an existing document.\nJust go ahead start clicking screenshots and notes button. It will append to your document.')
                    window['session'].update(session_value)
            except:
                sg.popup('Alert!!','Alert!! Folder you are trying to create is already Present. Try giving a new Folder Name.\nRecommended to give a date as postfix. eg: Client_Reporting_20-08-20. \nIf you want to append to an existing document.\nJust go ahead start clicking screenshots and notes button. It will append to your document.')
                
        if event == "Take Screenshot":
            val=values['path']+'/'+values['session']
            img,string=take_screenshot(val)
            try:
                if i_count==1:
                    if n_taken == False:
                        if os.path.isfile(val+'/'+'mydoc.docx'):
                                pass
                        else:
                            mydoc = docx.Document()
                            mydoc.save(val+'/'+"mydoc"+'.docx')
                    add_image_to_doc(mydoc,img,val,string)
                    print("Screenshot {}".format(i_count)," taken.")
                    i_count+=1
                    ss_taken=True
                else:
                    add_image_to_doc(mydoc,img,val,string)
                    print("Screenshot {}".format(i_count)," taken.")
                    i_count+=1
                    ss_taken=True
            except:
                sg.popup('Alert!!','Make Sure your Word Document is Closed.')
                
        
        if event == "Take Notes":
            val=values['path']+'/'+values['session']
            content=values['notes']
            try:
                if t_count==1:
                    try:
                        val=values['path']+'/'+values['session']
                        os.makedirs(values['path']+'/'+values['session'])
                        os.chdir(values['path']+'/'+values['session'])
                        if ss_taken==False:
                            if os.path.isfile(val+'/'+'mydoc.docx'):
                                pass
                            else:
                                mydoc = docx.Document()
                                mydoc.save(val+'/'+"mydoc"+'.docx')
                        add_content_to_doc(mydoc,content,val)
                        window['notes'].update('')
                        print("Notes {}".format(t_count)," taken.")
                        t_count+=1
                        n_taken=True
                    except:
                        
                        add_content_to_doc(mydoc,content,val)
                        window['notes'].update('')
                        print("Notes {}".format(t_count)," taken.")
                        t_count+=1
                        n_taken=True
                    
                    
                else:
                    add_content_to_doc(mydoc,content,val)
                    window['notes'].update('')
                    print("Notes {}".format(t_count)," taken.")
                    t_count+=1
                    n_taken=True
            except:
                sg.popup('Alert!!  Make sure you closed the word document before you add notes!!')
        if event == sg.WIN_CLOSED:
            break
except:
    print("Some error")
    sys.exit()
os.chdir(og_path)
window.close()